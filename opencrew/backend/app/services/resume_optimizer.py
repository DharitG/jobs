import subprocess
import uuid
import logging
import tempfile
from pathlib import Path
from fastapi import UploadFile, HTTPException
import shutil
import os
from docx import Document
import spacy
from typing import List, Optional, Dict, Any
from sklearn.feature_extraction.text import TfidfVectorizer
import re # For basic text cleaning
from openai import OpenAI, OpenAIError # Added for Phase 4
import os # To get API key

# Import the schemas defined for structured output
from app.schemas.optimize import (
    ParsedResume, ResumeSection, ResumeItem,
    KeywordAnalysisRequest, KeywordAnalysisResponse, MissingTerm,
    RewriteRequest, RewriteResponse,
    Patch, ApplyPatchRequest, ApplyPatchResponse # Added patch schemas
)
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.units import inch

logger = logging.getLogger(__name__)

# Define the directory to store converted docx files
# In a real app, this might be cloud storage or a persistent volume
CONVERTED_DOCX_DIR = Path(tempfile.gettempdir()) / "resume_optimizer_docx"
CONVERTED_DOCX_DIR.mkdir(parents=True, exist_ok=True)

# Ensure unoconv is available
if not shutil.which("unoconv"):
    logger.warning(
        "unoconv command not found. Resume ingestion will likely fail. "
        "Please install unoconv and ensure LibreOffice is available."
    )

# --- Phase 2: Parse & Model ---

# Load spaCy model (consider loading only once at app startup for efficiency)
# You might need to download the model first: python -m spacy download en_core_web_sm
try:
    nlp = spacy.load("en_core_web_sm")
    logger.info("spaCy model 'en_core_web_sm' loaded.")
except OSError:
    logger.warning(
        "spaCy model 'en_core_web_sm' not found. "
        "Please download it: python -m spacy download en_core_web_sm. "
        "Parsing features might be limited."
    )
    nlp = None # Set to None if loading fails

# --- Helper Functions for Parsing ---

def is_likely_heading(paragraph) -> bool:
    """
    Basic heuristic to guess if a paragraph is a section heading.
    Checks for common heading styles, all caps, or specific keywords.
    This needs significant refinement for real-world resumes.
    """
    text = paragraph.text.strip()
    if not text:
        return False

    # Check for common heading styles (exact names depend on template)
    # Common examples: 'Heading 1', 'Heading 2', 'Title'
    if paragraph.style and paragraph.style.name.lower().startswith('heading'):
        return True

    # Check if text is all uppercase (and relatively short)
    if text.isupper() and len(text.split()) < 5:
        return True

    # Check for bold text (if the entire paragraph is bold)
    # Note: This checks the first run only, more robust check needed for mixed formatting
    if paragraph.runs and all(run.bold for run in paragraph.runs if run.text.strip()):
         # Check if it's not excessively long (headings are usually concise)
        if len(text.split()) < 8:
            return True


    # TODO: Add more sophisticated checks (e.g., font size changes, specific keywords)
    # common_headings = ["experience", "education", "skills", "summary", "objective", "projects"]
    # if text.lower() in common_headings:
    #     return True

    return False

def get_paragraph_type(paragraph) -> str:
    """Determines if a paragraph is likely a bullet point or standard text."""
    text = paragraph.text.strip()
    # Check style names (common in Word templates)
    if paragraph.style and paragraph.style.name.lower().contains('list paragraph'):
        return "bullet"
    # Check for common bullet characters (simplistic)
    if text.startswith(("*", "-", "•", "▪", "▫")): # Add more bullet types if needed
        return "bullet"
    # TODO: Check paragraph formatting properties (indentation, bullet formatting) via lxml if needed
    return "paragraph"


# --- Core Parsing Function ---

def parse_docx_to_structure(docx_path: Path) -> ParsedResume:
    """
    Parses a DOCX file into a structured JSON-like format (using Pydantic models).

    Args:
        docx_path: Path to the DOCX file.

    Returns:
        A ParsedResume object representing the structured content.

    Raises:
        FileNotFoundError: If the docx_path does not exist.
        Exception: For errors during DOCX parsing.
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found at: {docx_path}")

    try:
        document = Document(docx_path)
        parsed_resume = ParsedResume(sections=[])
        current_section = ResumeSection(title=None, items=[])

        for i, paragraph in enumerate(document.paragraphs):
            paragraph_text = paragraph.text.strip()
            if not paragraph_text: # Skip empty paragraphs
                continue

            if is_likely_heading(paragraph):
                # If we have items in the current section, add it to the resume
                if current_section.items:
                    parsed_resume.sections.append(current_section)
                # Start a new section
                current_section = ResumeSection(title=paragraph_text, items=[])
                logger.debug(f"Detected heading: '{paragraph_text}'")
            else:
                # Add item to the current section
                item_type = get_paragraph_type(paragraph)
                current_section.items.append(
                    ResumeItem(type=item_type, content=paragraph_text)
                )

        # Add the last section if it has items
        if current_section.items or current_section.title:
             # If the very first thing was content without a heading
            if not parsed_resume.sections and not current_section.title:
                 current_section.title = "General" # Assign a default title
            parsed_resume.sections.append(current_section)


        # TODO: Implement table parsing if needed
        # for table in document.tables:
        #     # Process table content...
        #     pass

        logger.info(f"Successfully parsed DOCX '{docx_path.name}' into {len(parsed_resume.sections)} sections.")
        return parsed_resume

    except Exception as e:
        logger.exception(f"Error parsing DOCX file {docx_path}: {e}")
        # Re-raise or handle appropriately
        raise Exception(f"Failed to parse DOCX file: {e}")


# --- Phase 3: Keyword Gap Analysis ---

def clean_text(text: str) -> str:
    """Basic text cleaning: lowercase, remove punctuation."""
    text = text.lower()
    text = re.sub(r'[^\w\s]', '', text) # Remove punctuation
    return text

def analyze_keyword_gaps(
    request: KeywordAnalysisRequest
) -> KeywordAnalysisResponse:
    """
    Performs keyword gap analysis between a job description and a parsed resume
    using TF-IDF.

    Args:
        request: KeywordAnalysisRequest containing job description and parsed resume.

    Returns:
        KeywordAnalysisResponse containing a list of potentially missing terms.
    """
    logger.info("Starting keyword gap analysis.")

    # 1. Extract text from parsed resume
    resume_text = ""
    for section in request.parsed_resume.sections:
        if section.title:
            resume_text += section.title + " "
        for item in section.items:
            resume_text += item.content + " "
    resume_text = clean_text(resume_text)
    logger.debug(f"Resume text length (cleaned): {len(resume_text)}")

    # 2. Clean job description text
    jd_text = clean_text(request.job_description)
    logger.debug(f"Job Description text length (cleaned): {len(jd_text)}")


    # 3. Use TF-IDF to find important terms in the Job Description
    # We fit TF-IDF only on the JD to find terms important *to the JD*.
    # We could also fit on both JD and resume corpus for relative importance,
    # but this approach focuses on JD requirements first.
    if not jd_text.strip():
         logger.warning("Job description text is empty after cleaning.")
         return KeywordAnalysisResponse(missing_terms=[])

    try:
        vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1, 3)) # Use 1-3 word phrases
        tfidf_matrix = vectorizer.fit_transform([jd_text])
        feature_names = vectorizer.get_feature_names_out()
        scores = tfidf_matrix.toarray().flatten()

        # Get terms sorted by TF-IDF score
        jd_terms_scores = sorted(zip(feature_names, scores), key=lambda x: x[1], reverse=True)

        # Filter terms with non-zero score
        jd_important_terms = {term: score for term, score in jd_terms_scores if score > 0.05} # Use a threshold
        logger.debug(f"Found {len(jd_important_terms)} important terms in JD (TF-IDF > 0.05).")

    except ValueError as e:
        # Can happen if jd_text is empty or only contains stop words
        logger.warning(f"TF-IDF Vectorizer failed: {e}. Likely due to empty/stopword-only JD.")
        return KeywordAnalysisResponse(missing_terms=[])


    # 4. Identify terms present in JD but potentially missing/rare in Resume
    missing_terms_list: List[MissingTerm] = []
    resume_words = set(resume_text.split()) # Simple check for presence

    # More robust check: Check for phrase presence
    for term, score in jd_important_terms.items():
        # Simple check: if the term (as a whole phrase) isn't in the resume text
        if term not in resume_text:
             # Check if individual words are present (less strict)
             # words_in_term = term.split()
             # if not all(word in resume_words for word in words_in_term):
             missing_terms_list.append(MissingTerm(term=term, score=round(score, 4)))


    # Sort missing terms by score (most important first)
    missing_terms_list.sort(key=lambda x: x.score, reverse=True)

    # Limit the number of suggestions?
    max_suggestions = 20
    missing_terms_list = missing_terms_list[:max_suggestions]

    logger.info(f"Keyword analysis complete. Found {len(missing_terms_list)} potential missing terms.")

    return KeywordAnalysisResponse(missing_terms=missing_terms_list)


# --- Phase 4: Rewrite Engine ---

# Initialize OpenAI client
# Ensure OPENAI_API_KEY environment variable is set
try:
    openai_client = OpenAI()
    # Test connection (optional, but good practice)
    # openai_client.models.list()
    logger.info("OpenAI client initialized successfully.")
except OpenAIError as e:
    logger.error(f"Failed to initialize OpenAI client: {e}. Ensure OPENAI_API_KEY is set.")
    openai_client = None # Set to None to prevent usage if init fails

def generate_rewrite_suggestion(request: RewriteRequest) -> RewriteResponse:
    """
    Uses an LLM (GPT via OpenAI API) to rewrite a resume item, injecting keywords.

    Args:
        request: RewriteRequest containing the original item and terms to inject.

    Returns:
        RewriteResponse containing the suggested rewritten content.

    Raises:
        HTTPException: If the OpenAI client is not available or the API call fails.
    """
    if not openai_client:
        logger.error("OpenAI client is not available.")
        raise HTTPException(status_code=500, detail="Rewrite engine (OpenAI) is not configured.")

    logger.info(f"Generating rewrite suggestion for item: '{request.original_item.content[:50]}...'")

    # Construct the prompt for the LLM
    missing_terms_str = ", ".join([f"'{term}'" for term in request.missing_terms])
    original_text = request.original_item.content
    item_type = request.original_item.type # e.g., "bullet", "paragraph"

    # Basic prompt - needs refinement based on testing
    prompt = f"""Rewrite the following resume {item_type} to naturally incorporate the keywords/phrases: {missing_terms_str}.

Constraints:
- Maintain the original meaning and tone.
- Keep the length similar to the original (ideally within +/- 15% characters).
- If it's a bullet point, keep it as a concise bullet point starting with an action verb if appropriate.
- If it's a paragraph, keep it as a paragraph.
- Ensure grammatical correctness and professional language.
- Do NOT add any introductory phrases like "Here's the rewritten version:". Output only the rewritten text.

Original {item_type}:
"{original_text}"

Rewritten {item_type}:
"""
    if request.context:
         prompt = f"Context: This text is from the '{request.context}' section of a resume.\n\n" + prompt

    logger.debug(f"Generated LLM Prompt:\n{prompt}")

    try:
        # Use the chat completions endpoint (recommended)
        completion = openai_client.chat.completions.create(
            model="gpt-4-turbo-preview", # Or another suitable model like gpt-3.5-turbo
            messages=[
                {"role": "system", "content": "You are an expert resume editor."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5, # Lower temperature for more focused output
            max_tokens=len(original_text.split()) + 50, # Estimate max tokens needed
            n=1, # Generate one suggestion
            stop=None # Let the model decide when to stop
        )

        rewritten_content = completion.choices[0].message.content.strip()

        # Basic cleanup: remove potential quotation marks if the model added them
        if rewritten_content.startswith('"') and rewritten_content.endswith('"'):
            rewritten_content = rewritten_content[1:-1]

        logger.info("Successfully received rewrite suggestion from LLM.")
        logger.debug(f"Rewritten content: {rewritten_content}")

        return RewriteResponse(rewritten_content=rewritten_content)

    except OpenAIError as e:
        logger.error(f"OpenAI API error during rewrite: {e}")
        raise HTTPException(status_code=502, detail=f"Failed to get rewrite suggestion from AI: {e}")
    except Exception as e:
        logger.exception(f"Unexpected error during rewrite generation: {e}")
        raise HTTPException(status_code=500, detail="An unexpected error occurred during rewrite generation.")


# --- Phase 5: Apply Patches ---

def apply_patches_to_docx(request: ApplyPatchRequest) -> Path:
    """
    Applies text replacement patches to a DOCX file.

    Finds paragraphs matching the original_content of each patch and replaces
    their text with rewritten_content, attempting to preserve formatting.

    Args:
        request: ApplyPatchRequest containing the original DOCX path and patches.

    Returns:
        Path to the newly created, modified DOCX file.

    Raises:
        FileNotFoundError: If the original DOCX file is not found.
        Exception: For errors during DOCX processing or saving.
    """
    original_path = Path(request.original_docx_path)
    if not original_path.exists():
        raise FileNotFoundError(f"Original DOCX file not found at: {original_path}")

    logger.info(f"Applying {len(request.patches)} patches to {original_path.name}")

    try:
        document = Document(original_path)
        patches_applied_count = 0

        # Create a mapping for quick lookup
        patch_map = {patch.original_content.strip(): patch.rewritten_content for patch in request.patches}

        for paragraph in document.paragraphs:
            original_para_text = paragraph.text.strip()

            if original_para_text in patch_map:
                new_text = patch_map[original_para_text]
                logger.debug(f"Found match for patch: '{original_para_text[:50]}...' -> '{new_text[:50]}...'")

                # Strategy: Clear existing runs, add a new run with the new text,
                # preserving the paragraph's style. This is simpler but might lose
                # intra-paragraph formatting (like a single bold word).
                # A more complex approach would involve finding the exact runs to modify.

                # Clear existing runs in the paragraph
                for run in paragraph.runs:
                    run.clear() # Clears the run content

                # Add a new run with the rewritten text. It should inherit paragraph style.
                # If the original had specific run formatting (like bold), it might be lost here.
                # We could try copying the style from the first original run if available.
                if paragraph.runs: # Check if runs existed before clearing
                     # This is heuristic: assumes first run's style is representative
                    style_run = paragraph.runs[0]
                    paragraph.add_run(new_text, style=style_run.style)
                    # Try to copy font properties (might not capture everything)
                    # new_run = paragraph.add_run(new_text)
                    # new_run.font.name = style_run.font.name
                    # new_run.font.size = style_run.font.size
                    # new_run.bold = style_run.bold
                    # new_run.italic = style_run.italic
                    # new_run.underline = style_run.underline
                else:
                     # If paragraph had no runs initially (unlikely for non-empty text)
                     paragraph.add_run(new_text)


                patches_applied_count += 1
                # Optional: Remove the applied patch from the map to avoid reapplying if text appears multiple times
                # del patch_map[original_para_text]
                # if not patch_map: break # Stop if all patches are applied

        # TODO: Add logic to patch text within tables if needed

        if patches_applied_count != len(request.patches):
             logger.warning(f"Applied {patches_applied_count} patches, but {len(request.patches)} were provided. Some original content might not have been found.")

        # Save the modified document to a new file
        modified_filename = f"{original_path.stem}_modified_{uuid.uuid4().hex[:8]}.docx"
        modified_path = original_path.parent / modified_filename
        document.save(modified_path)

        logger.info(f"Successfully applied patches. Modified DOCX saved to: {modified_path}")
        return modified_path

    except Exception as e:
        logger.exception(f"Error applying patches to DOCX file {original_path}: {e}")
        raise Exception(f"Failed to apply patches to DOCX file: {e}")


# --- Phase 6: Export ---

# Define directory for exported PDFs
EXPORTED_PDF_DIR = Path(tempfile.gettempdir()) / "resume_optimizer_pdfs"
EXPORTED_PDF_DIR.mkdir(parents=True, exist_ok=True)

def export_designer_pdf(docx_path: Path) -> Path:
    """
    Exports the DOCX file to a high-fidelity PDF using unoconv.

    Args:
        docx_path: Path to the input DOCX file (usually the modified one).

    Returns:
        Path to the created designer-friendly PDF file.

    Raises:
        FileNotFoundError: If the input DOCX file is not found.
        HTTPException: If unoconv is not found or conversion fails.
        Exception: For other unexpected errors.
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"Input DOCX file not found at: {docx_path}")
    if not shutil.which("unoconv"):
        logger.error("unoconv command not found for PDF export.")
        raise HTTPException(status_code=500, detail="File conversion tool (unoconv) not available for PDF export.")

    output_filename = f"{docx_path.stem}_designer.pdf"
    output_path = EXPORTED_PDF_DIR / output_filename

    logger.info(f"Attempting designer PDF export: {docx_path} -> {output_path}")

    try:
        command = ["unoconv", "-f", "pdf", "-o", str(output_path), str(docx_path)]
        process = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=60
        )

        if process.returncode != 0:
            logger.error(f"unoconv PDF export failed (code {process.returncode}): {process.stderr}")
            if output_path.exists():
                output_path.unlink()
            raise HTTPException(status_code=500, detail=f"Designer PDF conversion failed: {process.stderr}")

        if not output_path.exists() or output_path.stat().st_size == 0:
             logger.error(f"unoconv PDF export ran but output file is missing or empty: {output_path}")
             raise HTTPException(status_code=500, detail="Designer PDF conversion resulted in an empty file.")

        logger.info(f"Successfully exported designer PDF to: {output_path}")
        return output_path

    except subprocess.TimeoutExpired:
        logger.error(f"unoconv PDF export command timed out for file: {docx_path.name}")
        if output_path.exists():
            output_path.unlink()
        raise HTTPException(status_code=500, detail="Designer PDF conversion timed out.")
    except Exception as e:
        logger.exception(f"Error during designer PDF export for {docx_path.name}: {e}")
        if output_path.exists():
            output_path.unlink()
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during designer PDF export: {str(e)}")


def export_ats_pdf(docx_path: Path) -> Path:
    """
    Exports the DOCX content to a simple, text-only, ATS-friendly PDF using reportlab.

    Args:
        docx_path: Path to the input DOCX file.

    Returns:
        Path to the created ATS-friendly PDF file.

    Raises:
        FileNotFoundError: If the input DOCX file is not found.
        Exception: For errors during DOCX reading or PDF generation.
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"Input DOCX file not found at: {docx_path}")

    output_filename = f"{docx_path.stem}_ats_friendly.pdf"
    output_path = EXPORTED_PDF_DIR / output_filename

    logger.info(f"Attempting ATS-friendly PDF export: {docx_path} -> {output_path}")

    try:
        document = Document(docx_path)
        styles = getSampleStyleSheet()
        # Customize styles slightly if needed
        styles['Normal'].alignment = TA_LEFT
        styles['Heading1'].alignment = TA_LEFT # Assuming headings were parsed correctly

        story = []
        # Basic extraction - iterates paragraphs, ignores tables/images
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Rudimentary style mapping (can be improved based on parsing phase)
            style_name = 'Normal'
            if paragraph.style and paragraph.style.name.lower().startswith('heading'):
                 style_name = 'h1' # Map to reportlab Heading1 (needs definition) or use bold Normal
                 # Use a bold Normal style for simplicity unless specific heading styles are defined
                 style = styles['Normal']
                 # Create a paragraph with bold - requires HTML-like tags
                 # This might need refinement based on how reportlab handles inline styles
                 story.append(Paragraph(f"<b>{text}</b>", style))
                 story.append(Spacer(1, 0.1*inch)) # Add space after heading
            # elif paragraph.style and paragraph.style.name.lower().contains('list paragraph'):
            #     # Add indentation for bullets?
            #     style = styles['Normal']
            #     story.append(Paragraph(f"• {text}", style)) # Add bullet manually
            else:
                style = styles['Normal']
                story.append(Paragraph(text, style))
            story.append(Spacer(1, 0.05*inch)) # Small space between paragraphs

        # Build the PDF
        pdf_doc = SimpleDocTemplate(str(output_path))
        pdf_doc.build(story)

        logger.info(f"Successfully exported ATS-friendly PDF to: {output_path}")
        return output_path

    except Exception as e:
        logger.exception(f"Error during ATS-friendly PDF export for {docx_path.name}: {e}")
        if output_path.exists():
            output_path.unlink() # Clean up partial file
        raise Exception(f"An unexpected error occurred during ATS-friendly PDF export: {str(e)}")


# --- Phase 1: Ingest ---
# (convert_to_docx function remains below)

async def convert_to_docx(file: UploadFile) -> Path:
    """
    Converts an uploaded file (PDF, DOC, etc.) to DOCX format using unoconv.

    Args:
        file: The uploaded file object from FastAPI.

    Returns:
        The Path object pointing to the created DOCX file.

    Raises:
        HTTPException: If the conversion fails or unoconv is not found.
    """
    if not shutil.which("unoconv"):
        logger.error("unoconv command not found.")
        raise HTTPException(status_code=500, detail="File conversion tool (unoconv) not available.")

    # Create temporary files for input and output
    input_suffix = Path(file.filename).suffix if file.filename else ".tmp"
    output_filename = f"{uuid.uuid4()}.docx"
    output_path = CONVERTED_DOCX_DIR / output_filename

    try:
        # Use NamedTemporaryFile to handle cleanup automatically if possible
        with tempfile.NamedTemporaryFile(delete=False, suffix=input_suffix) as temp_input_file:
            # Write uploaded content to the temporary input file
            content = await file.read()
            temp_input_file.write(content)
            temp_input_path = Path(temp_input_file.name)

        logger.info(f"Attempting conversion: {temp_input_path} -> {output_path}")

        # Construct the unoconv command
        # Use --stdout and redirect to handle potential filename issues and ensure output goes where we want
        # Using a specific output path directly might be more reliable if stdout redirection causes issues
        # command = ["unoconv", "-f", "docx", "--stdout", str(temp_input_path)] # Option 1: stdout
        command = ["unoconv", "-f", "docx", "-o", str(output_path), str(temp_input_path)] # Option 2: direct output

        # Execute the command
        process = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=60  # Add a timeout (e.g., 60 seconds)
        )

        # Check for errors
        if process.returncode != 0:
            logger.error(f"unoconv failed (code {process.returncode}): {process.stderr}")
            # Attempt cleanup of potentially created (but empty/invalid) output file
            if output_path.exists():
                output_path.unlink()
            raise HTTPException(status_code=500, detail=f"File conversion failed: {process.stderr}")

        # Option 1: If using stdout, write the output to the file
        # with open(output_path, "w") as f_out:
        #     f_out.write(process.stdout)

        # Verify output file was created (especially important for Option 2)
        if not output_path.exists() or output_path.stat().st_size == 0:
             logger.error(f"unoconv ran but output file is missing or empty: {output_path}")
             raise HTTPException(status_code=500, detail="File conversion resulted in an empty file.")

        logger.info(f"Successfully converted '{file.filename}' to '{output_path}'")
        return output_path

    except subprocess.TimeoutExpired:
        logger.error(f"unoconv command timed out for file: {file.filename}")
        if output_path.exists():
            output_path.unlink() # Clean up partial file
        raise HTTPException(status_code=500, detail="File conversion timed out.")
    except Exception as e:
        logger.exception(f"Error during file conversion for {file.filename}: {e}")
        if output_path.exists():
            output_path.unlink() # Clean up partial file
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during file conversion: {str(e)}")
    finally:
        # Clean up the temporary input file
        if 'temp_input_path' in locals() and temp_input_path.exists():
            try:
                temp_input_path.unlink()
            except OSError as e:
                logger.warning(f"Could not delete temporary input file {temp_input_path}: {e}")
